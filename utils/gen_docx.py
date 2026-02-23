from docx import Document
from docx.shared import Pt
import os

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Титульный
doc.add_heading('Дорожная Карта и План Внедрения', 0)
doc.add_paragraph('Версия: 1.0 (Server-Ready)')

files = [
    r"c:\Users\андрей\.gemini\antigravity\brain\6c4f5a23-e97a-4bb2-a678-458471c7480a\task.md",
    r"c:\Users\андрей\.gemini\antigravity\brain\6c4f5a23-e97a-4bb2-a678-458471c7480a\implementation_plan.md"
]

for fpath in files:
    if os.path.exists(fpath):
        with open(fpath, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith("# "):
                    doc.add_heading(line.replace("# ", ""), level=1)
                elif line.startswith("## "):
                    doc.add_heading(line.replace("## ", ""), level=2)
                elif line.startswith("### "):
                    doc.add_heading(line.replace("### ", ""), level=3)
                elif line.startswith("* "):
                    p = doc.add_paragraph(line.replace("* ", ""), style='List Bullet')
                elif line.startswith("[ ]"):
                    p = doc.add_paragraph(line.replace("[ ]", "□ "), style='List Bullet')
                else:
                    doc.add_paragraph(line)

output_path = r"c:\Users\андрей\OneDrive\Desktop\ПРОЕКТЫ\fRieNDLee_FTP\Note\Roadmap_RUS.docx"
doc.save(output_path)
print(f"SUCCESS: {output_path}")
